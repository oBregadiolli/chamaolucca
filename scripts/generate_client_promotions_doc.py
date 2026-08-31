from pathlib import Path
from datetime import date

from PIL import Image, ImageDraw, ImageFont, ImageOps
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(r"C:\Users\brega\Projetos\ChamaoLucca")
OUT_DIR = ROOT / "docs" / "cliente"
ASSET_DIR = OUT_DIR / "assets_promocoes"
DOCX_PATH = OUT_DIR / "Entrega_Modulo_Promocoes_ChamaoLucca_Andrey.docx"

PRINT_ADMIN = Path(r"C:\Users\brega\AppData\Local\Temp\codex-clipboard-f195b537-3561-4802-8828-41ab90039a47.png")
PRINT_MODAL = Path(r"C:\Users\brega\AppData\Local\Temp\codex-clipboard-c3254236-63cf-4637-8503-e19e7fd1b4b3.png")

GREEN = "16A34A"
DARK_GREEN = "15803D"
LIGHT_GREEN = "F0FDF4"
MID_GREEN = "DCFCE7"
INK = "111827"
MUTED = "64748B"
SOFT = "F8FAFC"
BORDER = "DADFE8"
GOLD = "F59E0B"
RED = "EF4444"


def ensure_dirs():
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    ASSET_DIR.mkdir(parents=True, exist_ok=True)


def font(size=24, bold=False):
    candidates = [
        r"C:\Windows\Fonts\arialbd.ttf" if bold else r"C:\Windows\Fonts\arial.ttf",
        r"C:\Windows\Fonts\calibrib.ttf" if bold else r"C:\Windows\Fonts\calibri.ttf",
    ]
    for candidate in candidates:
        try:
            return ImageFont.truetype(candidate, size)
        except Exception:
            pass
    return ImageFont.load_default()


def rounded_rect(draw, xy, radius, fill, outline=None, width=1):
    draw.rounded_rectangle(xy, radius=radius, fill=fill, outline=outline, width=width)


def crop_assets():
    assets = {}

    if PRINT_ADMIN.exists():
      img = Image.open(PRINT_ADMIN).convert("RGB")
      # Recorta a área útil do navegador, removendo a conversa lateral e o toast inferior.
      crop = img.crop((495, 160, 1835, 875))
      crop = ImageOps.expand(crop, border=3, fill="#E2E8F0")
      path = ASSET_DIR / "painel_promocoes.png"
      crop.save(path)
      assets["painel"] = path

    if PRINT_MODAL.exists():
      img = Image.open(PRINT_MODAL).convert("RGB")
      # Mantém o modal e tira sobras pesadas da borda.
      crop = img.crop((90, 45, 855, 820))
      crop = ImageOps.expand(crop, border=3, fill="#E2E8F0")
      path = ASSET_DIR / "modal_promocao.png"
      crop.save(path)
      assets["modal"] = path

    flow = Image.new("RGB", (1400, 620), "#FFFFFF")
    d = ImageDraw.Draw(flow)
    title = font(42, True)
    h = font(29, True)
    body = font(24)
    small = font(20)

    d.text((54, 44), "Exemplo prático: brinde automático no carrinho", fill=f"#{INK}", font=title)
    d.text((56, 100), "O cliente não precisa lembrar da promoção. O sistema mostra o benefício sozinho.", fill=f"#{MUTED}", font=body)

    cards = [
        ("1", "Cliente adiciona café", "Café Maratá entra na sacola.", "☕"),
        ("2", "Promoção é liberada", "A regra identifica que o café dá direito ao leite.", "✓"),
        ("3", "Leite aparece grátis", "O leite entra como brinde automático e sai do carrinho se a regra deixar de valer.", "🎁"),
    ]
    x = 58
    for number, card_title, desc, icon in cards:
        rounded_rect(d, (x, 180, x + 390, 510), 28, "#F8FAFC", "#DADFE8", 3)
        rounded_rect(d, (x + 28, 210, x + 88, 270), 30, "#DCFCE7", "#BBF7D0", 2)
        d.text((x + 48, 225), number, fill=f"#{DARK_GREEN}", font=h, anchor="mm")
        d.text((x + 118, 214), card_title, fill=f"#{INK}", font=h)
        d.text((x + 32, 300), icon, fill=f"#{GREEN}", font=font(54, True))
        d.multiline_text((x + 32, 380), desc, fill=f"#{MUTED}", font=small, spacing=7)
        if number != "3":
            d.text((x + 410, 318), "→", fill=f"#{GREEN}", font=font(56, True))
        x += 455

    path = ASSET_DIR / "fluxo_brinde_automatico.png"
    flow.save(path)
    assets["fluxo"] = path

    return assets


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, color=BORDER):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = "w:" + edge
        el = borders.find(qn(tag))
        if el is None:
            el = OxmlElement(tag)
            borders.append(el)
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "8")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color)


def set_cell_margins(cell, top=120, start=160, bottom=120, end=160):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    mar = tc_pr.first_child_found_in("w:tcMar")
    if mar is None:
        mar = OxmlElement("w:tcMar")
        tc_pr.append(mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            cell.width = Inches(width)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            set_cell_border(cell)


def set_run(run, size=11, color=INK, bold=False, italic=False):
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    run.bold = bold
    run.italic = italic


def add_text(paragraph, text, size=11, color=INK, bold=False, italic=False):
    run = paragraph.add_run(text)
    set_run(run, size=size, color=color, bold=bold, italic=italic)
    return run


def add_para(doc, text="", size=11, color=INK, bold=False, italic=False, after=6, before=0, align=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.line_spacing = 1.18
    if align is not None:
        p.alignment = align
    add_text(p, text, size=size, color=color, bold=bold, italic=italic)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(16 if level == 1 else 10)
    p.paragraph_format.space_after = Pt(7 if level == 1 else 5)
    add_text(p, text, size=17 if level == 1 else 13, color=DARK_GREEN if level == 1 else INK, bold=True)
    return p


def add_bullet(doc, text, color=INK):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.line_spacing = 1.15
    p.runs.clear()
    add_text(p, text, size=10.5, color=color)
    return p


def add_callout(doc, title, body, fill=LIGHT_GREEN, accent=DARK_GREEN):
    table = doc.add_table(rows=1, cols=1)
    set_table_width(table, [6.3])
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    add_text(p, title, size=11.5, color=accent, bold=True)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    p2.paragraph_format.line_spacing = 1.15
    add_text(p2, body, size=10.5, color=INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return table


def add_status_table(doc):
    rows = [
        ("Ações da semana no painel", "Entregue", "O Andrey consegue criar campanhas sem depender de uma alteração específica para cada ideia."),
        ("Preço especial em produto", "Entregue", "Ex.: leite em pó por R$ 3,99 acima de R$ 20."),
        ("Produto grátis", "Entregue", "Ex.: comprou café, ganhou leite."),
        ("Brinde automático no carrinho", "Entregue", "O benefício aparece sozinho para o cliente quando a regra é liberada."),
        ("Datas e ativação", "Entregue", "Dá para ligar/desligar e definir período da ação."),
        ("Limite por pedido", "Entregue", "Ajuda a proteger a margem e evitar abuso da promoção."),
        ("Relatórios detalhados e campanhas por comportamento", "Próxima evolução", "Pode entrar depois, quando fizer sentido medir resultado e criar ações mais avançadas."),
    ]
    table = doc.add_table(rows=1, cols=3)
    set_table_width(table, [2.0, 1.35, 3.0])
    headers = ["Entrega", "Status", "O que isso significa"]
    for i, h in enumerate(headers):
        cell = table.cell(0, i)
        set_cell_shading(cell, SOFT)
        p = cell.paragraphs[0]
        add_text(p, h, size=10.5, color=INK, bold=True)
    for item, status, meaning in rows:
        cells = table.add_row().cells
        values = [item, status, meaning]
        for i, value in enumerate(values):
            set_cell_border(cells[i])
            set_cell_margins(cells[i])
            if status == "Entregue":
                set_cell_shading(cells[1], MID_GREEN)
            elif status == "Próxima evolução":
                set_cell_shading(cells[1], "FEF3C7")
            p = cells[i].paragraphs[0]
            add_text(p, value, size=9.8, color=DARK_GREEN if i == 1 and status == "Entregue" else INK, bold=i in (0, 1))
    doc.add_paragraph().paragraph_format.space_after = Pt(4)


def add_examples(doc):
    examples = [
        ("Leite em pó da semana", "Compra acima de R$ 20", "Leite em pó sai por R$ 3,99", "Ideal para chamar atenção e aumentar o valor da sacola."),
        ("Café com leite", "Cliente coloca café no carrinho", "Leite aparece como brinde automático", "Boa para combos e para dar sensação de vantagem imediata."),
        ("Açúcar promocional", "Compra acima de R$ 50", "1 açúcar sai por R$ 1,00", "Boa para campanha de giro e incentivo de ticket maior."),
    ]
    table = doc.add_table(rows=1, cols=4)
    set_table_width(table, [1.5, 1.7, 1.75, 1.4])
    for i, h in enumerate(["Campanha", "Quando libera", "O que acontece", "Por que é útil"]):
        cell = table.cell(0, i)
        set_cell_shading(cell, SOFT)
        add_text(cell.paragraphs[0], h, size=9.8, color=INK, bold=True)
    for row in examples:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_border(cells[i])
            set_cell_margins(cells[i])
            p = cells[i].paragraphs[0]
            add_text(p, value, size=9.2, color=INK, bold=i == 0)


def add_image(doc, path, caption, width=6.15):
    if not path or not Path(path).exists():
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(path), width=Inches(width))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(10)
    add_text(cap, caption, size=9, color=MUTED, italic=True)


def build_doc(assets):
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    add_text(header, "ChamaoLucca | Módulo de Promoções", size=9, color=MUTED)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_text(footer, "Documento de apresentação para Andrey", size=9, color=MUTED)

    add_para(doc, "CHAMAOLUCCA", size=10, color=GREEN, bold=True, after=4)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    add_text(p, "Módulo de Promoções Inteligentes", size=28, color=INK, bold=True)
    add_para(doc, "Uma entrega para criar ações da semana, combos e brindes automáticos diretamente pelo painel.", size=13, color=MUTED, after=16)
    add_callout(
        doc,
        "Resumo para o Andrey",
        "Foi criado um módulo para transformar promoções em regras configuráveis. Na prática, o ChamaoLucca passa a permitir campanhas como “leite em pó por R$ 3,99 acima de R$ 20” ou “comprou café, ganhou leite”, sem precisar criar cada ação manualmente do zero.",
    )
    add_para(doc, f"Preparado em {date.today().strftime('%d/%m/%Y')} para apresentação da entrega.", size=9.5, color=MUTED, italic=True, after=14)

    add_heading(doc, "O que foi feito", 1)
    add_para(doc, "Foi entregue uma área nova no painel administrativo para cadastrar promoções de forma simples. A ideia é permitir que o mercado monte ações comerciais com mais liberdade: ações da semana, produtos chamariz, brindes e campanhas para aumentar o valor médio dos pedidos.", after=8)
    add_status_table(doc)

    add_heading(doc, "Como funciona na prática", 1)
    add_para(doc, "O Andrey cria uma promoção preenchendo três ideias simples:", after=5)
    add_bullet(doc, "Quando a promoção deve liberar: por exemplo, compra acima de R$ 20 ou produto específico no carrinho.")
    add_bullet(doc, "O que o cliente ganha: preço especial em um produto ou produto grátis.")
    add_bullet(doc, "Quais limites protegem a margem: quantidade máxima por pedido, período da ação e se a promoção está ativa.")

    add_callout(
        doc,
        "O ponto forte da entrega",
        "O cliente não precisa saber a regra de cabeça. O carrinho passa a mostrar quando a oferta foi liberada e, no caso de brinde, o produto grátis aparece automaticamente na sacola.",
        fill="FFF7ED",
        accent="C2410C",
    )

    add_heading(doc, "Exemplos de campanhas que já fazem sentido", 1)
    add_examples(doc)

    add_heading(doc, "Prints da área de promoções", 1)
    add_para(doc, "Abaixo estão os principais pontos da experiência no painel: uma área exclusiva para acompanhar promoções e um modal para cadastrar novas ações de forma guiada.", after=8)
    add_image(doc, assets.get("painel"), "Painel administrativo com a nova área de Promoções.", width=6.35)
    add_image(doc, assets.get("modal"), "Tela de criação de promoção, onde a ação da semana é configurada.", width=5.65)

    add_heading(doc, "Brinde automático: a experiência que mais vende a ideia", 1)
    add_para(doc, "Esse comportamento foi incluído para deixar a promoção mais clara para o cliente final. Se a campanha diz que o cliente ganha um produto ao comprar outro, o brinde aparece sozinho no carrinho quando a regra é cumprida.", after=8)
    add_image(doc, assets.get("fluxo"), "Fluxo simples do brinde automático no carrinho.", width=6.35)
    add_bullet(doc, "Se o cliente adiciona o produto que libera a promoção, o brinde aparece como grátis.")
    add_bullet(doc, "Se o cliente remove o produto que liberava a promoção, o brinde some automaticamente.")
    add_bullet(doc, "Na revisão antes do pagamento, o brinde também aparece, deixando a vantagem visível e transparente.")

    add_heading(doc, "Por que isso tem valor para o ChamaoLucca", 1)
    add_callout(
        doc,
        "Não é só uma promoção fixa",
        "A entrega foi pensada como uma base reutilizável. Em vez de criar uma campanha isolada para açúcar, café ou leite, o painel permite montar novas ações comerciais conforme a necessidade da semana.",
    )
    add_bullet(doc, "Ajuda a aumentar o valor médio das compras, porque o cliente enxerga quanto falta para liberar uma vantagem.")
    add_bullet(doc, "Facilita ações rápidas de mercado, como queima de estoque, produto da semana e combos.")
    add_bullet(doc, "Dá mais autonomia para o time criar campanhas sem depender de uma nova programação a cada ideia.")
    add_bullet(doc, "Deixa a experiência mais clara para o cliente, porque o benefício aparece dentro da sacola e na revisão do pedido.")

    add_heading(doc, "Possibilidades futuras", 1)
    add_para(doc, "A base criada já abre caminho para evoluções maiores. Esses itens não precisam entrar agora para a entrega fazer sentido; são caminhos naturais para quando o ChamaoLucca quiser campanhas mais avançadas.", after=6)
    add_bullet(doc, "Frete grátis por primeira compra ou por valor mínimo.")
    add_bullet(doc, "Campanhas para clientes recorrentes ou clientes que não compram há alguns dias.")
    add_bullet(doc, "Limite total da campanha, como “válido para os 100 primeiros pedidos”.")
    add_bullet(doc, "Relatórios mostrando quantas vezes a promoção foi usada e quanto ajudou a vender.")
    add_bullet(doc, "Regras de acumulação, para decidir quando uma promoção pode ou não somar com outra.")

    add_heading(doc, "Mensagem final", 1)
    add_para(doc, "Essa entrega deixa o ChamaoLucca com uma ferramenta comercial mais flexível. A loja passa a ter uma forma prática de criar ações da semana, destacar produtos estratégicos e transformar o carrinho em um espaço de incentivo à compra, com uma experiência mais clara para o cliente e mais controle para a operação.", after=8)
    add_para(doc, "Foi uma construção com bastante cuidado porque envolve painel, carrinho, revisão do pedido e comportamento automático da promoção. O resultado é uma base que resolve as campanhas atuais e prepara o caminho para ações comerciais mais fortes no futuro.", size=11.3, color=INK, bold=True, after=0)

    doc.save(DOCX_PATH)


if __name__ == "__main__":
    ensure_dirs()
    assets = crop_assets()
    build_doc(assets)
    print(DOCX_PATH)
